"""Structured data inspection + query tools used by the answer MCP server.

Three groups of helpers, all task-local (operate only on a given task_dir):

1. inspect_data(task_dir) -> dict
   Scan context/ once, report a compact schema summary of every CSV / JSON /
   SQLite / PDF / DOCX file. The summary is the most cost-effective thing the
   agent can read before deciding which file to query and how.

2. sqlite_query(db_path, sql) -> dict
   Safe, read-only SELECT against a SQLite database file. Results capped at
   MAX_ROWS rows and MAX_CELLS total cells.

3. pandas_query(csv_path, expression) -> dict
   Evaluate a pandas expression on a single CSV. The expression operates on a
   variable named `df`, which is `pd.read_csv(csv_path)`. We sandbox `eval`
   to a small whitelist of names. Result must be a Series, DataFrame, or
   scalar; we serialize the head.

All three never raise to the caller; they always return a dict whose
`is_error` flag and `message` describe failures, so the LLM can recover.
"""
from __future__ import annotations

import csv
import json
import sqlite3
from pathlib import Path
from typing import Any

DEFAULT_MAX_ROWS = 200
DEFAULT_MAX_CELLS = 6000
DEFAULT_SAMPLE_VALUES = 8
INT_LIKE_LIMIT = 50  # if column unique <= this, list all values
MAX_FILE_BYTES_FOR_FULL_SCAN = 50 * 1024 * 1024  # 50 MB

# Optional deps (loaded lazily so this module imports cleanly even if missing)
try:
    import pandas as pd  # noqa: F401
except Exception:  # pragma: no cover
    pd = None  # type: ignore

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None  # type: ignore

try:
    import pdfplumber as _pdfplumber
except Exception:  # pragma: no cover
    _pdfplumber = None  # type: ignore

try:
    import docx  # python-docx
except Exception:  # pragma: no cover
    docx = None  # type: ignore


def _short_text(text: str, limit: int = 240) -> str:
    text = (text or "").strip()
    if len(text) <= limit:
        return text
    return text[:limit] + "…"


def _csv_summary(path: Path, max_sample_values: int = DEFAULT_SAMPLE_VALUES) -> dict[str, Any]:
    info: dict[str, Any] = {"path": str(path), "kind": "csv"}
    try:
        size_bytes = path.stat().st_size
        info["size_bytes"] = size_bytes
        with path.open("r", encoding="utf-8-sig", newline="", errors="replace") as handle:
            reader = csv.reader(handle)
            try:
                header = next(reader)
            except StopIteration:
                info["columns"] = []
                info["row_count"] = 0
                return info

            columns = [c.strip() for c in header]
            n_cols = len(columns)
            counts: list[int] = [0] * n_cols
            uniques: list[set[str]] = [set() for _ in range(n_cols)]
            samples: list[list[str]] = [[] for _ in range(n_cols)]
            row_count = 0
            for row in reader:
                row_count += 1
                for i in range(min(n_cols, len(row))):
                    val = row[i].strip()
                    if val == "":
                        continue
                    counts[i] += 1
                    if len(uniques[i]) < 1000:
                        uniques[i].add(val)
                    if len(samples[i]) < max_sample_values and val not in samples[i]:
                        samples[i].append(val)

            cols_summary = []
            for i, name in enumerate(columns):
                unique_set = uniques[i]
                unique_count = len(unique_set)
                values_field: list[str] | None = None
                if unique_count <= INT_LIKE_LIMIT:
                    values_field = sorted(unique_set)
                # dtype guess
                non_empty = [v for v in samples[i] if v != ""]
                dtype_guess = "string"
                if non_empty:
                    if all(v.lstrip("-").isdigit() for v in non_empty):
                        dtype_guess = "int"
                    else:
                        try:
                            for v in non_empty:
                                float(v)
                            dtype_guess = "float"
                        except ValueError:
                            dtype_guess = "string"
                cols_summary.append({
                    "name": name,
                    "non_empty_count": counts[i],
                    "unique_count_capped_1000": unique_count,
                    "dtype_guess": dtype_guess,
                    "samples": samples[i],
                    "values": values_field,  # only present for low-cardinality columns
                })
            info["columns"] = cols_summary
            info["row_count"] = row_count
    except Exception as exc:  # pragma: no cover
        info["is_error"] = True
        info["message"] = f"failed to summarize csv: {exc!r}"
    return info


def _json_summary(path: Path, max_sample_values: int = DEFAULT_SAMPLE_VALUES) -> dict[str, Any]:
    info: dict[str, Any] = {"path": str(path), "kind": "json"}
    try:
        info["size_bytes"] = path.stat().st_size
        with path.open("r", encoding="utf-8", errors="replace") as handle:
            data = json.load(handle)

        def describe(node: Any, depth: int = 0) -> dict[str, Any]:
            if isinstance(node, dict):
                keys = list(node.keys())
                desc = {"type": "object", "keys": keys[:30]}
                if depth < 2 and keys:
                    sample_key = keys[0]
                    desc["sample_value"] = describe(node[sample_key], depth + 1)
                return desc
            if isinstance(node, list):
                desc = {"type": "array", "length": len(node)}
                if node and depth < 2:
                    desc["sample_item"] = describe(node[0], depth + 1)
                return desc
            return {"type": type(node).__name__, "value": _short_text(str(node), 80)}

        info["structure"] = describe(data, 0)

        # Heuristic: if the top-level is a list of records, summarize as a "table".
        records: list[dict[str, Any]] | None = None
        if isinstance(data, list) and data and isinstance(data[0], dict):
            records = data
        elif isinstance(data, dict):
            for k, v in data.items():
                if isinstance(v, list) and v and isinstance(v[0], dict):
                    records = v
                    info["records_field"] = k
                    break
        if records is not None:
            info["row_count"] = len(records)
            field_keys: dict[str, Any] = {}
            for rec in records[:1000]:
                for fk, fv in rec.items():
                    if fk not in field_keys:
                        field_keys[fk] = {"non_empty_count": 0, "samples": [], "uniques": set()}
                    if fv not in (None, ""):
                        field_keys[fk]["non_empty_count"] += 1
                        if len(field_keys[fk]["samples"]) < max_sample_values:
                            sval = _short_text(str(fv), 80)
                            if sval not in field_keys[fk]["samples"]:
                                field_keys[fk]["samples"].append(sval)
                        if len(field_keys[fk]["uniques"]) < 1000:
                            field_keys[fk]["uniques"].add(str(fv))
            cols = []
            for fk, agg in field_keys.items():
                unique_count = len(agg["uniques"])
                values_field = sorted(agg["uniques"]) if unique_count <= INT_LIKE_LIMIT else None
                cols.append({
                    "name": fk,
                    "non_empty_count": agg["non_empty_count"],
                    "unique_count_capped_1000": unique_count,
                    "samples": agg["samples"],
                    "values": values_field,
                })
            info["columns"] = cols
    except Exception as exc:
        info["is_error"] = True
        info["message"] = f"failed to summarize json: {exc!r}"
    return info


def _sqlite_summary(path: Path) -> dict[str, Any]:
    info: dict[str, Any] = {"path": str(path), "kind": "sqlite"}
    try:
        info["size_bytes"] = path.stat().st_size
        # readonly via URI
        uri = f"file:{path.as_posix()}?mode=ro"
        conn = sqlite3.connect(uri, uri=True)
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name")
            tables = [row[0] for row in cursor.fetchall()]
            tables_summary = []
            for tname in tables:
                cursor.execute(f"PRAGMA table_info('{tname}')")
                cols = [
                    {"name": row[1], "dtype": row[2], "pk": bool(row[5])}
                    for row in cursor.fetchall()
                ]
                cursor.execute(f"SELECT COUNT(*) FROM '{tname}'")
                rcount = cursor.fetchone()[0]
                tables_summary.append({
                    "name": tname,
                    "columns": cols,
                    "row_count": rcount,
                })
            info["tables"] = tables_summary
        finally:
            conn.close()
    except Exception as exc:
        info["is_error"] = True
        info["message"] = f"failed to summarize sqlite: {exc!r}"
    return info


def _pdf_summary(path: Path) -> dict[str, Any]:
    info: dict[str, Any] = {"path": str(path), "kind": "pdf"}
    try:
        info["size_bytes"] = path.stat().st_size
        # Try pdfplumber first for table detection
        if _pdfplumber is not None:
            with _pdfplumber.open(str(path)) as doc:
                info["page_count"] = len(doc.pages)
                if doc.pages:
                    p0 = doc.pages[0]
                    text = p0.extract_text() or ""
                    info["page1_excerpt"] = _short_text(text, 600)
                    # Count tables across first 3 pages
                    table_count = 0
                    for pg in doc.pages[:3]:
                        table_count += len(pg.extract_tables() or [])
                    if table_count:
                        info["table_count_first3pages"] = table_count
        elif fitz is not None:
            with fitz.open(str(path)) as doc:
                info["page_count"] = doc.page_count
                if doc.page_count > 0:
                    first_text = doc[0].get_text()
                    info["page1_excerpt"] = _short_text(first_text, 600)
        else:
            info["is_error"] = True
            info["message"] = "neither pdfplumber nor pymupdf installed"
    except Exception as exc:
        info["is_error"] = True
        info["message"] = f"failed to summarize pdf: {exc!r}"
    return info


def _docx_summary(path: Path) -> dict[str, Any]:
    info: dict[str, Any] = {"path": str(path), "kind": "docx"}
    if docx is None:
        info["is_error"] = True
        info["message"] = "python-docx not installed"
        return info
    try:
        info["size_bytes"] = path.stat().st_size
        document = docx.Document(str(path))
        paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        info["paragraph_count"] = len(paragraphs)
        info["first_paragraphs"] = [_short_text(p, 200) for p in paragraphs[:8]]
        # tables count
        info["table_count"] = len(document.tables)
    except Exception as exc:
        info["is_error"] = True
        info["message"] = f"failed to summarize docx: {exc!r}"
    return info


def inspect_data(task_dir: Path, max_files_per_kind: int = 50) -> dict[str, Any]:
    """Scan context/ and return per-file schema summaries.

    Markdown / text files are intentionally skipped because there is a separate
    MarkdownRagIndex tool for those.
    """
    context_dir = task_dir / "context"
    out: dict[str, Any] = {
        "task_dir": str(task_dir),
        "context_dir": str(context_dir),
        "csv": [],
        "json": [],
        "sqlite": [],
        "pdf": [],
        "docx": [],
        "video": [],
        "skipped": [],
    }
    if not context_dir.exists():
        out["error"] = "context dir does not exist"
        return out

    for path in sorted(context_dir.rglob("*")):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        try:
            if suffix == ".csv":
                if len(out["csv"]) < max_files_per_kind:
                    out["csv"].append(_csv_summary(path))
            elif suffix == ".json":
                if len(out["json"]) < max_files_per_kind:
                    out["json"].append(_json_summary(path))
            elif suffix in {".sqlite", ".db", ".sqlite3"}:
                if len(out["sqlite"]) < max_files_per_kind:
                    out["sqlite"].append(_sqlite_summary(path))
            elif suffix == ".pdf":
                if len(out["pdf"]) < max_files_per_kind:
                    out["pdf"].append(_pdf_summary(path))
            elif suffix == ".docx":
                if len(out["docx"]) < max_files_per_kind:
                    out["docx"].append(_docx_summary(path))
            elif suffix in {".mp4", ".avi", ".mov", ".mkv", ".webm"}:
                if len(out["video"]) < max_files_per_kind:
                    out["video"].append({"path": str(path), "kind": "video",
                                         "size_bytes": path.stat().st_size, "ext": suffix})
            else:
                out["skipped"].append({"path": str(path), "ext": suffix})
        except Exception as exc:  # pragma: no cover
            out["skipped"].append({"path": str(path), "ext": suffix, "error": repr(exc)})
    return out


# ----- query helpers --------------------------------------------------


def sqlite_query(db_path: Path, sql: str, max_rows: int = DEFAULT_MAX_ROWS) -> dict[str, Any]:
    sql_stripped = sql.strip().rstrip(";")
    if not sql_stripped:
        return {"is_error": True, "message": "sql is empty"}
    lowered = sql_stripped.lower()
    if not (lowered.startswith("select") or lowered.startswith("with")):
        return {"is_error": True, "message": "only SELECT / WITH queries are allowed"}
    forbidden = ("attach ", "detach ", "pragma ", "insert ", "update ", "delete ",
                 "drop ", "create ", "alter ", "replace ", "vacuum")
    if any(token in lowered for token in forbidden):
        return {"is_error": True, "message": "query contains a non-read keyword"}
    try:
        uri = f"file:{db_path.as_posix()}?mode=ro"
        conn = sqlite3.connect(uri, uri=True)
        try:
            cursor = conn.cursor()
            cursor.execute(sql_stripped)
            cols = [c[0] for c in (cursor.description or [])]
            rows = cursor.fetchmany(max_rows)
            return {
                "columns": cols,
                "rows": [list(r) for r in rows],
                "row_count": len(rows),
                "truncated": len(rows) >= max_rows,
            }
        finally:
            conn.close()
    except Exception as exc:
        return {"is_error": True, "message": f"sqlite error: {exc!r}"}


_PANDAS_SAFE_NAMES: dict[str, Any] = {}


def _pandas_safe_globals() -> dict[str, Any]:
    if _PANDAS_SAFE_NAMES:
        return _PANDAS_SAFE_NAMES
    if pd is None:
        return {}
    _PANDAS_SAFE_NAMES.update({
        "pd": pd,
        "len": len,
        "min": min,
        "max": max,
        "sum": sum,
        "abs": abs,
        "round": round,
        "sorted": sorted,
        "set": set,
        "list": list,
        "dict": dict,
        "tuple": tuple,
        "range": range,
    })
    return _PANDAS_SAFE_NAMES


def pandas_query(
    csv_path: Path,
    expression: str,
    max_rows: int = DEFAULT_MAX_ROWS,
    extra_csvs: dict[str, Path] | None = None,
) -> dict[str, Any]:
    if pd is None:
        return {"is_error": True, "message": "pandas is not installed"}
    if not csv_path.exists():
        return {"is_error": True, "message": f"csv not found: {csv_path}"}
    try:
        df = pd.read_csv(csv_path)
    except Exception as exc:
        return {"is_error": True, "message": f"failed to read csv: {exc!r}"}

    locals_dict: dict[str, Any] = {"df": df}
    if extra_csvs:
        for name, p in extra_csvs.items():
            try:
                locals_dict[name] = pd.read_csv(p)
            except Exception as exc:
                return {"is_error": True, "message": f"failed to read extra csv {name}: {exc!r}"}

    globals_dict = _pandas_safe_globals()

    expr = expression.strip()
    if not expr:
        return {"is_error": True, "message": "expression is empty"}
    if "__" in expr or "import" in expr or "open(" in expr or "exec(" in expr:
        return {"is_error": True, "message": "expression contains forbidden tokens"}

    try:
        result = eval(expr, {"__builtins__": {}, **globals_dict}, locals_dict)  # noqa: S307
    except Exception as exc:
        return {"is_error": True, "message": f"expression failed: {exc!r}"}

    return _serialize_pandas_result(result, max_rows=max_rows)


def _serialize_pandas_result(result: Any, max_rows: int) -> dict[str, Any]:
    if pd is None:
        return {"is_error": True, "message": "pandas not available"}
    if isinstance(result, pd.DataFrame):
        head = result.head(max_rows)
        return {
            "kind": "DataFrame",
            "columns": [str(c) for c in head.columns],
            "rows": head.astype(object).where(pd.notna(head), None).values.tolist(),
            "row_count_returned": int(len(head)),
            "row_count_total": int(len(result)),
            "truncated": len(result) > max_rows,
        }
    if isinstance(result, pd.Series):
        head = result.head(max_rows)
        return {
            "kind": "Series",
            "name": str(result.name) if result.name is not None else None,
            "values": head.astype(object).where(pd.notna(head), None).tolist(),
            "row_count_returned": int(len(head)),
            "row_count_total": int(len(result)),
            "truncated": len(result) > max_rows,
        }
    if isinstance(result, (int, float)):
        return {"kind": "scalar", "value": result}
    return {"kind": type(result).__name__, "value": _short_text(repr(result), 800)}


def read_pdf_pages(pdf_path: Path, page_start: int = 1, page_end: int = 1, extract_tables: bool = False) -> dict[str, Any]:
    if _pdfplumber is None and fitz is None:
        return {"is_error": True, "message": "neither pdfplumber nor pymupdf installed"}
    if not pdf_path.exists():
        return {"is_error": True, "message": f"pdf not found: {pdf_path}"}
    try:
        # Use pdfplumber for better table extraction when available
        if _pdfplumber is not None:
            with _pdfplumber.open(str(pdf_path)) as doc:
                page_count = len(doc.pages)
                page_start = max(1, page_start)
                page_end = min(page_count, max(page_start, page_end))
                chunks = []
                for i in range(page_start - 1, page_end):
                    page = doc.pages[i]
                    text = page.extract_text() or ""
                    page_data: dict[str, Any] = {"page": i + 1, "text": text}
                    # Extract tables if requested
                    if extract_tables:
                        tables = page.extract_tables()
                        if tables:
                            page_data["tables"] = tables
                    chunks.append(page_data)
                return {
                    "page_count": page_count,
                    "pages": chunks,
                    "page_start": page_start,
                    "page_end": page_end,
                }
        else:
            # Fallback to pymupdf (text only)
            with fitz.open(str(pdf_path)) as doc:
                page_start = max(1, page_start)
                page_end = min(doc.page_count, max(page_start, page_end))
                chunks = []
                for i in range(page_start - 1, page_end):
                    chunks.append({"page": i + 1, "text": doc[i].get_text()})
                return {
                    "page_count": doc.page_count,
                    "pages": chunks,
                    "page_start": page_start,
                    "page_end": page_end,
                }
    except Exception as exc:
        return {"is_error": True, "message": f"pdf read failed: {exc!r}"}


def read_docx_full(docx_path: Path, max_paragraphs: int = 200) -> dict[str, Any]:
    if docx is None:
        return {"is_error": True, "message": "python-docx not installed"}
    if not docx_path.exists():
        return {"is_error": True, "message": f"docx not found: {docx_path}"}
    try:
        document = docx.Document(str(docx_path))
        paragraphs = []
        for p in document.paragraphs:
            text = p.text
            if text and text.strip():
                paragraphs.append({"style": getattr(p.style, "name", ""), "text": text})
            if len(paragraphs) >= max_paragraphs:
                break
        return {
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
            "table_count": len(document.tables),
        }
    except Exception as exc:
        return {"is_error": True, "message": f"docx read failed: {exc!r}"}
